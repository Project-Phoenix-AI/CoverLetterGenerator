from docx import Document
from docx.shared import Pt
from sys import argv,exit
from docx2pdf import convert
import os
from utils import get_company_name, get_hiring_manager_name
import argparse

parser = argparse.ArgumentParser()
parser.add_argument('--pos','-p',help = 'Shortcut for the position name',metavar='cs',type = str,default = 'ds')
parser.add_argument('--company','-c',help = 'Company name',metavar = 'meta',type = str,nargs = "+")
parser.add_argument('--manager','-m',help ='The name of the hiring manager',metavar = 'john',type = str,default = 'Dear Hiring Manager,',nargs = "+")
args = parser.parse_args()

interest = {'intern':'internship','cs':'software engineer','ml':'machine learning engineer','ds':'data scientist', 'deng':'data engineer', 'mlops':'MLOps engineer','da':'data analyst'}
positions = {'intern':'computer vision','cs':'software engineering','ml':'machine learning engineering','ds':'data science','deng':'data engineering','mlops':'MLOps engineering','da':'data science'}

# if len(argv) >= 3:
#     company = get_company_name(argv)
#     hiring_manager = get_hiring_manager_name(argv)
# else:
#     print(f"usage: python cover.py {list(positions.keys())} 'company name'")
#     exit(0)
position = args.pos

if isinstance(args.manager,list):
    hiring_manager = ' '.join(list(map(lambda x: x.capitalize(),args.manager)))
    hiring_manager = f"Dear {hiring_manager},"
else:
    hiring_manager = args.manager

company = ' '.join(args.company)

if len(company) <= 1:
    print('Enter the company',company)
    exit(1)

if position not in positions:
    print('invalid position:', position, )
    print('valid positions:', list(positions.keys()))
    exit(1)

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Calibri'#'Times New Roman'#'Calibri'
font.size = Pt(11)

greetings = hiring_manager
greetings = document.add_paragraph(greetings)

body = f'''
I am writing for {interest[position]} position at {company}! I am pursuing an MSc in Machine Learning at KTH Royal Institute of Technology. I have a bachelor's degree in degree in Electronics and Communication Engineering from Istanbul Technical University with a GPA 3.42 out of 4.0.  

I worked full-time as a Data Scientist while doing my master's, illustrating my ability to manage tight deadlines effectively. I have experience in google cloud tools. In my work, I have implemented data processing workflows using google BigQuery, Cloud Functions, Cloud Run and Pub/Sub. I developed neural networks and machine learning models for forecasting battery lifetime.'''

body = document.add_paragraph(body)
body.alignment = 0  # '3' corresponds to 'justify' alignment in docx

bullet_point_entry = document.add_paragraph('Throughout my career, I have developed expertise in:')

bullet_points = {'ml':[
                    'Google Cloud Platform (GCP), Google Big Query, Cloud Functions, Vertex AI, Qlik Sense',
                    'pandas, scikit-learn, keras, pytorch, tensorflow, flask, seaborn, matplotlib',
                    'Github Actions, Terraform, Docker'],
                'ds':[
                    'Google Cloud Platform (GCP), Google Big Query, Cloud Functions, Vertex AI, Qlik Sense',
                    'pandas, scikit-learn, keras, tensorflow, flask, seaborn, matplotlib',
                    'Github Actions, Terraform, Docker'],
                'cs':[
                    'Google Cloud Platform (GCP), Google Big Query, Cloud Functions, Vertex AI, Qlik Sense',
                    'pandas, scikit-learn, keras, tensorflow, flask, seaborn, matplotlib',
                    'Github Actions, Terraform, Docker'],
                'deng':[
                    'Google Cloud Platform (GCP), Google Big Query, Cloud Functions, Vertex AI, Qlik Sense',
                    'pandas, scikit-learn, keras, tensorflow, flask, seaborn, matplotlib',
                    'Github Actions, Terraform, Docker'],
                'all':[
                    'Google Cloud, Google BigQuery, Cloud Functions, Pub/Sub, Qlik Sense',
                    'Github Actions, Terraform, Docker',
                    'Pandas, Scikit-learn, Keras, Pytorch, Tensorflow, Flask, Seaborn, Matplotlib'
                    ],
                }

for point in bullet_points['all']:
    bullet = document.add_paragraph(point)
    bullet.style = 'List Bullet'


summary =f'My experience and academic background make me a great candidate for this position. Thank you for considering my application. I am thrilled to discuss how my skills and experiences might assist in the success of your team.'
summmary = document.add_paragraph(summary)


greetings = document.add_paragraph('Sincerely,')


contact = document.add_paragraph('Berkan Yapıcı\nPhone: +46729161069 | Email: berkanyapici9956@gmail.com')

company_name = company.replace(' ','_')
doc_name = f'cover_letter_{company_name}_{position}.docx'


if not os.path.exists('cover_letters/docs'):
    os.mkdir('cover_letters/docs')

path = f'cover_letters/docs/{doc_name}'
document.save(path)

# if not os.path.exists('cover_letters/pdfs'):
#     os.mkdir('cover_letters/pdfs')
    
# convert(path, f'cover_letters/pdfs/{doc_name.replace("docx","pdf")}')

print(f'cover letter generated\ncompany:{company}\nposition:{positions[position].capitalize()}')





